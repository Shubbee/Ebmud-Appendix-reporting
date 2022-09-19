"""
Task: Read Excel files, create tables, graphs, and insert them into a MS Word document
Created on %(date)s

@author: Shubhneet Singh 
ssin@dhigroup.com
DHI,US
"""

import os
import time
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from datetime import date
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

task = 'Task: Read Excel files, create tables, graphs, and insert them into a MS Word document\n'
day = date.today().strftime("%B%d, %Y")
tool_starttime = time.time()

#Task directory
wdir = r"C:\Users\ssin\OneDrive - DHI\Desktop\EBMUD\\"
os.chdir(wdir)

# Data files:
appendixE_exceldir = r".\AppendixE\\"  # AppendixE
appendixF_exceldir = r".\AppendixF\\"  # AppendixE       
raincsv_loc = r".\FY22_GARR_ITA_Dataset.xlsx"
raincsv_sheetname = 'EBMUD_Basin-ITA_Dataset_2021-10'
datatable_path = r"Table_Template.xlsx"
templatetable_path = r"Table_Template.docx"

# Create log for comments, assumption and notes:
log_file = open("Readme_PeakUpdatedFormula.txt","w+")
log_file.write('{} \nDeveloped by Shubhneet Singh\nssin@dhigroup.com\n{}\n\nNotes and Assumptions:\n'.format(task, day))
 
         
#%% Tool

log_file.write('\n For peak recorded and error calculations, Column C was used') 

# Start/End times: constraint reading, plotting data (include limits)
plot_starttime = pd.to_datetime('2021-11-1')
plot_endtime = pd.to_datetime('2022-4-15')
log_file.write('\nPlots time bounds:\nStart Time: {}\nEnd Time: {}'.format(plot_starttime, plot_endtime)) 

#Rain Data
rain_df = pd.read_excel(raincsv_loc,
                        sheet_name = raincsv_sheetname,
                        index_col= 'Time(PST)')
log_file.write('\nRainfall data source {}'.format(raincsv_loc)) 

tablenumber = 0
appendix_doc = Document()
templatetable = Document(templatetable_path).tables[0]._tbl #From template word doc
xlsheets = ['Measured Data FY22', 'PICS_Flow', 'Rain', 'VOLUME', 'Scatter Input Data']

# for appendix in ['E', 'F']:
for appendix in ['E']:   
    
    appendix_dir = r".\Appendix{}\\".format(appendix)
    appendix_xlfilenames = os.listdir(appendix_dir)  
    appendix_doc.add_heading('Appendix{}'.format(appendix), 1)
    datatable = pd.read_excel(datatable_path) #From template worksheet    
    log_file.write('\nAppendix{} excel files count - {}'.format(appendix, len(appendix_xlfilenames)))
    
    for xlf, xlfname in enumerate(appendix_xlfilenames):   
    # for xlf, xlfname in enumerate(appendix_xlfilenames[0:1]):
        
        xlfpath = appendix_dir + xlfname
        xlsheets_data = pd.read_excel(xlfpath,
                                      sheet_name = xlsheets) 
        
        meter_name = xlsheets_data['Measured Data FY22'].columns[1]
        if meter_name == '': print(xlfname + ': Missing meter name')
        ita_name = xlsheets_data['Scatter Input Data'].iloc[6,4]  
        if ita_name == '':   print(xlfname + ': Missing ita name')
    
        ## Plot the results from workbooks - one plot/workbook:        
        fig, Results = plt.subplots(figsize=(8.39,6.2))
        #Observed Data:
        observed_index = xlsheets_data['Measured Data FY22'].iloc[3:,0]
        observed_filter = (observed_index >= plot_starttime) & (observed_index <= plot_endtime)
        observed_x = observed_index[observed_filter]
        observed_y = xlsheets_data['Measured Data FY22'].iloc[3:,7][observed_filter]
        obs_ax = Results.plot(observed_x, observed_y,
                              color = 'black',
                              linewidth = 0.2,
                              label = 'Flow',
                              alpha = .8)
        #Modeled Data:
        modeled_index = xlsheets_data['PICS_Flow'].iloc[:,0]
        modeled_filter = (modeled_index >= plot_starttime) & (modeled_index <= plot_endtime)
        modeled_x = modeled_index[modeled_filter] 
        modeled_y = xlsheets_data['PICS_Flow'].iloc[:,1][modeled_filter]     
        mod_ax = Results.plot(modeled_x, modeled_y,
                              color = 'red',
                              linewidth = 0.2,
                              label = 'PICS Flow',
                              alpha = .8)
        #Rain data on secondary axis
        if 'and' in ita_name:
            rain_name = ita_name.split('and')[0][:-1]
            log_file.write('\nPlot {}: Rain picked from ita {}'.format(xlfname, rain_name)) 
        else:
            rain_name = ita_name
        secax = Results.twinx()
        rain_x = rain_df.index
        rain_y = rain_df.loc[:,rain_name]
        rain_ax = secax.plot(rain_x, rain_y,
                             color = 'blue',
                             label = 'Rain',
                             linewidth = 0.4,
                             alpha = .8)
        #Plot Format
        Results.set_title('ITA ' + ita_name + ': ' + meter_name, fontsize=8)
        # X-axis:    
        plt.gcf().autofmt_xdate()
        dtformat = mdates.DateFormatter('%m/%d/%Y')
        plt.gca().xaxis.set_major_formatter(dtformat)
        plt.gca().xaxis.set_major_locator(mdates.MonthLocator(bymonthday=[1,15]))
        Results.xaxis.set_tick_params(labelsize = 6)
        # Y-axis:
        rainmax = rain_y.max()
        y1max = np.nanmax([np.nanmax(observed_y),
                           np.nanmax(modeled_y)])
        
        Results.set_ylim(0, (1.2+rainmax)*y1max)                            
        Results.set_ylabel("Flow (MGD)", fontsize=8)        
        Results.xaxis.grid(linewidth=0.2)
        Results.yaxis.grid(linewidth=0.2)
        Results.yaxis.set_tick_params(labelsize = 7)

        # Y2-axis:
        secax.set_xlim(plot_starttime, plot_endtime)
        secax.set_ylim(1,0)
        secax.set_ylabel('Rain (inch)', color = 'blue', fontsize=8)
        secax.tick_params(axis = 'y', labelcolor = 'blue')
        secax.yaxis.set_tick_params(labelsize = 7)        
        # Legend:
        all_ax = obs_ax + mod_ax + rain_ax
        labels = [ax.get_label() for ax in all_ax]
        Results.legend(all_ax, labels,
                       loc = 2,
                       bbox_to_anchor=(0.84, 0.97),
                       fontsize=7)    
        # Save plot:
        plots_path = r".\Appendix{}-Plots\\".format(appendix)
        png_name = xlfname[:-5]
        png_path = plots_path +  png_name
        if not os.path.exists(plots_path):
            os.makedirs(plots_path)        
        plt.savefig(png_path,
                    bbox_inches = 'tight',
                    dpi = 300)
        plt.close(fig)  
        
        ## Compute summary table items:
        for r in range(len(datatable)):
            if datatable.iloc[r,1].isoformat() != 'NaT':
                
                start_t = datatable.iloc[r,1]
                end_t = datatable.iloc[r,2]
                               
                observedevent_filter = (observed_index >= start_t) & (observed_index <= end_t)
                observedevent = xlsheets_data['Measured Data FY22'].iloc[3:,7][observedevent_filter] #With updated formula                
                observedevent_Vol = observedevent.sum()/24/60*5     
                observedevent_original = xlsheets_data['Measured Data FY22'].iloc[3:,2][observedevent_filter]
                observedevent_originalVol = observedevent_original.sum()/24/60*5 
                
                modeledevent_filter = (modeled_index >= start_t) & (modeled_index <= end_t)
                modeledevent = xlsheets_data['PICS_Flow'].iloc[:,1][modeledevent_filter]
                modeledevent_Vol = modeledevent.sum()/24/60*15
                
                if observedevent_Vol != 0:                    
                    error = int(round(((modeledevent_Vol - observedevent_Vol)/ observedevent_Vol)*100))
                    observedevent_Vol = round(observedevent.sum()/24/60*5, 2)
                    # observedevent_peak = round(observedevent.max(), 2)
                    observedevent_peak = round(observedevent_original.max(), 2)  
                    if modeledevent_Vol != 0 and observedevent_originalVol != 0:
                        # error_peak = int(round(((modeledevent.max() - observedeventl.max())/ observedevent.max())*100))
                        error_peak = int(round(((modeledevent.max() - observedevent_original.max())/ observedevent_original.max())*100))
                    else:
                        error_peak = np.nan 
                        log_file.write('\nCheck {}: Modeled/Observed Volume 0 for event {}:{}'.format(xlfname, start_t, end_t))          
                else:
                    observedevent_Vol = np.nan
                    error = np.nan
                    observedevent_peak = np.nan
                    error_peak = np.nan     
                
                datatable.iloc[r,3] = observedevent_Vol
                datatable.iloc[r,4] = round(modeledevent_Vol, 2)
                datatable.iloc[r,5] = error                
                datatable.iloc[r,6] = observedevent_peak
                datatable.iloc[r,7] = round(modeledevent.max(), 2)
                datatable.iloc[r,8] = error_peak
                
        ## Plug plot, table  into a word table:
        new_table = deepcopy(templatetable) # Template table created before the loop 
        paragraph = appendix_doc.add_paragraph() 
        paragraph.paragraph_format.space_before = Pt(18)  
        paragraph._p.addnext(new_table)        
        worktable = appendix_doc.tables[tablenumber] 
        worktable.rows[0].cells[4].text= 'ITA ' + ita_name + ': ' + meter_name
        worktable.rows[0].cells[4].paragraphs[0].runs[0].font.bold = True
        worktable.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(9)
        worktable.rows[0].cells[4].paragraphs[0].alignment = 1
        #Add data to table
        updaterows = [r for r in range(2, len(datatable)+2) if r not in [16, 17]]
        updatecolumns = [r for r in range(3, len(datatable.columns))]
        for r in updaterows:
            for c in updatecolumns:
                celldata = datatable.iloc[r-2,c]
                if ((c in [5, 8]) & (not np.isnan(celldata))):
                    celldata = int(celldata)
                worktable.rows[r].cells[c].text = ['' if np.isnan(celldata) else str(celldata)][0]
                worktable.rows[r].cells[c].paragraphs[0].runs[0].font.size = Pt(9)
                worktable.rows[r].cells[c].paragraphs[0].paragraph_format.space_after = Pt(0)
                worktable.rows[r].cells[c].paragraphs[0].paragraph_format.space_before = Pt(0)
                worktable.rows[r].cells[c].paragraphs[0].alignment = 1
        #Add Plot        
        addplot = appendix_doc.add_picture(png_path + '.png',
                                           width = Inches(8.39))        
        last_paragraph = appendix_doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph.paragraph_format.space_before = Pt(20)    
        appendix_doc.add_page_break()
        tablenumber += 1
appendix_doc.save('Appendix2022_PeakUpdatedFormula.docx')

#%% Log file time entry

tool_endtime = time.time()
time_taken = str(round((tool_endtime - tool_starttime)/60,0))
print('\n\n############\n')
print('\nTime taken: {}'.format(time_taken) + ' minutes')

log_file.write('\n\n############\n')               
log_file.write('\nTime taken: {}'.format(time_taken) + ' minutes')
log_file.close()