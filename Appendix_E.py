"""
Created on Mon Oct 26 14:03:55 2020

Read Excel file, plot table and graphs in MS Word

@author: ssin
DHI, Denver

"""

from itertools import takewhile 
from os import listdir
from os.path import isfile, join
import xlrd
import datetime
import matplotlib.pyplot as plt
import matplotlib.dates as mdates 
from docx import Document
from copy import deepcopy
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
from docx.shared import Inches



# Excel file dir and workbook names:
    
excel_location = r"C:\Users\ssin\OneDrive - DHI\Desktop\EBMUD\1_PICS_Results_Comparison\\"  #Appendix E
excel_files = [f for f in listdir(excel_location) if isfile(join(excel_location, f))]

# Function to read the length of column in Excel worksheet

def column_len(sheet, index):
    col_values = sheet.col_values(index)
    col_len = len(col_values)
    for _ in takewhile(lambda x: not x if x!=0 else x , reversed(col_values)):
        col_len -= 1
    return col_len


### Read the relevant data from workbook

# for a in range(len(excel_files)):
for a in range(53,54):
        
    workbook = xlrd.open_workbook(excel_location + excel_files[a])
    
    meter_sheet = workbook.sheet_by_name('Measured Data FY20')
    pics_sheet  = workbook.sheet_by_name('PICS_Flow')
    rain_sheet  = workbook.sheet_by_name('Rain')
    volume_sheet= workbook.sheet_by_name('VOLUME')
    ID_sheet = workbook.sheet_by_name('Scatter Input Data')
    
    meter_x_len = column_len(meter_sheet,0)
    meter_y_len = column_len(meter_sheet,3)
    if meter_x_len > meter_y_len:
        meter_x_len = meter_y_len
        
    
        
    pics_x_len = column_len(pics_sheet,0)
    pics_y_len = column_len(pics_sheet,1)
   
    
    rain_x_len = column_len(rain_sheet,0)
    rain_y_len = column_len(rain_sheet,1)
    
    volume_row_len = column_len(volume_sheet,14)
    volume_col_len = 9
    
    meter_name = ID_sheet.cell_value(7, 2)
    ita = ID_sheet.cell_value(7, 4)
    
    if meter_name == '':
        plot_name= excel_files[a]
        meter_name= plot_name[15:28]
        
    if ita == '':
        plot_name= excel_files[a]
        ita = plot_name[9:14]
    

    
    
    # Meter time-series
    meter_xx = [datetime.datetime(*xlrd.xldate_as_tuple(meter_sheet.cell_value(b, 0), workbook.datemode)) for b in range(1,meter_x_len-1) if meter_sheet.cell_value(b,3) !='' ]
    meter_yy = [meter_sheet.cell_value(c, 3) for c in range(1,meter_y_len-1)if meter_sheet.cell_value(c, 3) !='']
    
    clip_meter_start=np.where([meter_xx[i]==datetime.datetime(2019, 11, 1, 0, 0) for i in range(len(meter_xx))])[0]
    clip_meter_end=np.where([meter_xx[i]==datetime.datetime(2020, 4, 15, 0, 0) for i in range(len(meter_xx))])[0]
    if len(clip_meter_start) ==0:
        clip_meter_start = [0]
    if len(clip_meter_end) ==0:
        clip_meter_end = [len(meter_xx)]
        
    meter_x= meter_xx[clip_meter_start[0]:clip_meter_end[0]]
    meter_y= meter_yy[clip_meter_start[0]:clip_meter_end[0]]
    
    # PICS time-series
    pics_xx = [datetime.datetime(*xlrd.xldate_as_tuple(pics_sheet.cell_value(d, 0), workbook.datemode)) for d in range(1,pics_x_len-1) if pics_sheet.cell_value(d, 1) !='']
    pics_yy = [pics_sheet.cell_value(e, 1) for e in range(1,pics_y_len-1) if pics_sheet.cell_value(e, 1) !='']  
    
    if len(pics_xx) < len(pics_yy):
        for dd in range(len(pics_xx),len(pics_yy)):
            pics_xx.append(pics_xx[dd-1] + (pics_xx[1]-pics_xx[0]))    
            
    clip_pics_start=np.where([pics_xx[i]==datetime.datetime(2019, 11, 1, 0, 0) for i in range(len(pics_xx))])[0]
    clip_pics_end=np.where([pics_xx[i]==datetime.datetime(2020, 4, 15, 0, 0) for i in range(len(pics_xx))])[0]
    
    pics_x= pics_xx[clip_pics_start[0]:clip_pics_end[0]]
    pics_y= pics_yy[clip_pics_start[0]:clip_pics_end[0]]
    
    
    # Rain time-series
    rain_x = [datetime.datetime(*xlrd.xldate_as_tuple(rain_sheet.cell_value(f, 0), workbook.datemode)) for f in range(1,rain_x_len-1) if rain_sheet.cell_value(f, 1) !='']
    rain_y = [rain_sheet.cell_value(g, 1) for g in range(1,rain_y_len-1) if rain_sheet.cell_value(g, 1) !='']
    

### Plot the results from workbooks - one plot/workbook:
    
    fig, Results = plt.subplots(constrained_layout=True, figsize=(8.39,6.2))
    p1=Results.plot(meter_x, meter_y, 'black', linewidth=0.5, label='Flow')
    p2=Results.plot(pics_x, pics_y, 'red', linewidth=0.5, label = 'PICS flow')
    
    Results.set_ylabel("Flow (MGD)")
    Results.set_title('ITA ' + ita + ': ' + meter_name )
    Results.xaxis.grid()
    Results.yaxis.grid()
    Results.autoscale()

    secax = Results.twinx()
    color = 'tab:blue'
    p3=secax.plot(rain_x, rain_y, color=color, label = 'Rain', linewidth=0.75)
    secax.set_ylabel('Rain (inch)', color=color)
    secax.tick_params(axis='y', labelcolor=color)
    secax.set_ylim(1,0)
    secax.set_xlim(datetime.datetime(2019, 11, 1, 0, 0),datetime.datetime(2020, 4, 15, 0, 0))
    
    # Legend:
    p_sum = p1+p2+p3
    labs = [l.get_label() for l in p_sum]
    Results.legend(p_sum, labs, loc=2, bbox_to_anchor=(0.8, 0.9)) 
    
    # Date axis formating:
    plt.gcf().autofmt_xdate()
    myFmt = mdates.DateFormatter('%m-%d-%Y')
    plt.gca().xaxis.set_major_formatter(myFmt)
       
    png_file = excel_files[a]
    plt.savefig(r"C:\Users\ssin\OneDrive - DHI\Desktop\EBMUD\3_Plots\\" +  png_file[:-5], bbox_inches='tight', dpi=200)
    #plt.savefig(r"C:\Users\ssin\OneDrive - DHI\Desktop\EBMUD\5_Plot_Calibration\\" +  png_file[:-5], bbox_inches='tight')
    plt.close(fig)  
    
 
### Read data from workbook and plug in word table: 
   
    # Data from workbook:    
    volume_cols = [0,1,2,9,10,11,13,14,15]
    volume_data=[] 
    volume_start=[] 
    volume_end=[] 
    volume_rel_rows= list(range(3,volume_row_len))
    volume_rel_rows.extend([45,46,47])             

 
    for row in volume_rel_rows:
        
        if volume_sheet.cell_value(row, 1) !='':            
            volume_start = (datetime.datetime.strftime(xlrd.xldate_as_datetime(volume_sheet.cell_value(row, 1), workbook.datemode),'%m/%d/%y %H:%M'))
            volume_end = (datetime.datetime.strftime(xlrd.xldate_as_datetime(volume_sheet.cell_value(row, 2), workbook.datemode),'%m/%d/%y %H:%M'))
        else:
            volume_start = ('')
            volume_end = ('')
            
        for col in volume_cols:
    
            if col == volume_cols[0]:
                volume_data.append(volume_sheet.cell_value(row, col)) 
                        
            if col == volume_cols[1]:
               volume_data.append(volume_start) 
               
            if col == volume_cols[2]:
               volume_data.append(volume_end) 
               
            if col == volume_cols[3] or col == volume_cols[4]: 
                if volume_sheet.cell_value(row, col) != '':
                    Volume_MG = volume_sheet.cell_value(row, col)/10**6
                    volume_data.append(round(Volume_MG,2))
                else:
                    volume_data.append('')    
            
            if col == volume_cols[5] or col == volume_cols[8]:            
                if volume_sheet.cell_value(row, col) != '':
                    volume_data.append(int(volume_sheet.cell_value(row, col)*100))
                else:
                    volume_data.append('')   
                    
            if col == volume_cols[6] or col == volume_cols[7]:             
                if volume_sheet.cell_value(row, col) != '':
                    volume_data.append(round(volume_sheet.cell_value(row, col),2))
                else:
                    volume_data.append('')    
    
### Report - word document:
    
    table_template_doc = Document('AppendixE_SampleTable.docx')
    # doc = Document('AppendixE_DHI.docx')
    doc = Document('Table_005.docx')
        
    template = table_template_doc.tables[0] 
    tb1 = template._tbl
    # Copy of the table
    new_table = deepcopy(tb1)    

    # Add plot
    appendix_plot = doc.add_picture(r"C:\Users\ssin\OneDrive - DHI\Desktop\EBMUD\3_Plots" + '\\' +png_file[:-5] + '.png', width=Inches(8.39))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format_plot = last_paragraph.paragraph_format
    paragraph_format_plot.space_before = Pt(16)
    
             
    # Add data to the template table
    doc.add_page_break()    
    paragraph = doc.add_paragraph()
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph._p.addnext(new_table)
    
    table = doc.tables[0] 
    table.rows[0].cells[4].text= 'ITA ' + ita + ': ' + meter_name
    table.rows[0].cells[4].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[4].paragraphs[0].runs[0].font.size = Pt(9)
    table.rows[0].cells[4].paragraphs[0].alignment = 1
    
    
    for h in range(2,volume_row_len+2):
        for i in range(volume_col_len):
            table.rows[h].cells[i].text = str(volume_data[9*(h-2)+i])
            table.rows[h].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            table.rows[h].cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            table.rows[h].cells[i].paragraphs[0].paragraph_format.space_before = Pt(0)
        for i in range(3,volume_col_len):    
            table.rows[h].cells[i].paragraphs[0].alignment = 1

    doc.save('Table_005.docx')
    # os.system('AppendixE_DHI.docx')
    #doc.save('AppendixF_DHI.docx')
    