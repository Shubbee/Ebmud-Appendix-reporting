"""
Created on Thu Nov 12 15:43:29 2020
Read Excel file, plot table and graphs in MS Word

@author: ssin
DHI, Denver
"""


from itertools import takewhile 
from os import listdir
from os.path import isfile, join
import xlrd
import datetime
import matplotlib.pyplot as plt
import matplotlib.dates as mdates 
from docx import Document
import numpy as np

# Excel file dir and workbook names:
    
excel_location = r"C:\Users\ssin\OneDrive - DHI\Desktop\EBMUD\9_AppendixA_QDV_Scattergraph\AppendixA_Materials\\"
excel_files = [f for f in listdir(excel_location) if isfile(join(excel_location, f))]

# Function to read the length of column in Excel worksheet

def column_len(sheet, index):
    col_values = sheet.col_values(index)
    col_len = len(col_values)
    for _ in takewhile(lambda x: not x if x!=0 else x , reversed(col_values)):
        col_len -= 1
    return col_len


### Read the relevant data from workbook

for a in range(len(excel_files)):
# for a in range(1):
        
    workbook = xlrd.open_workbook(excel_location + excel_files[a])
    
    meter_sheet = workbook.sheet_by_name('Measured Data FY20')
    ID_sheet = workbook.sheet_by_name('Scatter Input Data')
    geometory_sheet = workbook.sheet_by_name('Geometry Circular')
    mannings_sheet = workbook.sheet_by_name('Mannings')
    
    meter_x_len = column_len(meter_sheet,0)
    meter_y_len = column_len(meter_sheet,3)
    if meter_x_len > meter_y_len:
        meter_x_len = meter_y_len
        
    meter_name = meter_sheet.cell_value(0, 7)
    ita = meter_sheet.cell_value(1, 7)
    if a==2:
        ita = '02-3'
        meter_name = 'SSD_0095B_001_Rev01'
        

    # Meter time-series
    
    meter_x = [datetime.datetime(*xlrd.xldate_as_tuple(meter_sheet.cell_value(b, 0), workbook.datemode)) for b in range(1,meter_x_len-1) if meter_sheet.cell_value(b,3) !='' ]
    meter_q = [meter_sheet.cell_value(c, 3) for c in range(1,meter_y_len-1)if meter_sheet.cell_value(c, 3) !='']
    meter_d = [meter_sheet.cell_value(c, 1) for c in range(1,meter_y_len-1)if meter_sheet.cell_value(c, 3) !='']
    meter_v = [meter_sheet.cell_value(c, 2) for c in range(1,meter_y_len-1)if meter_sheet.cell_value(c, 3) !='']
    
    nov_start = np.where(meter_x == np.datetime64(datetime.datetime(2019, 11, 1, 0, 0)))
    dec_end = np.where(meter_x == np.datetime64(datetime.datetime(2019, 12, 31, 23,45)))
    
    if nov_start[0].size != 0 and dec_end[0].size != 0:
        meter_v_novdec = meter_v[nov_start[0][0]:dec_end[0][0]]
        meter_d_novdec = meter_d[nov_start[0][0]:dec_end[0][0]]
    if nov_start[0].size == 0 and dec_end[0].size != 0:
        meter_v_novdec = meter_v[0:dec_end[0][0]]
        meter_d_novdec = meter_d[0:dec_end[0][0]]
    if nov_start[0].size != 0 and dec_end[0].size == 0:
        meter_v_novdec = meter_v[nov_start[0][0]:len(meter_x)]
        meter_d_novdec = meter_d[nov_start[0][0]:len(meter_x)]
    if nov_start[0].size == 0 and dec_end[0].size == 0:
        meter_v_novdec = meter_v[0:len(meter_x)]
        meter_d_novdec = meter_d[0:len(meter_x)]


    crown_d = [ID_sheet.cell_value(c, 0) for c in range(43,45)]
    crown_v = [ID_sheet.cell_value(c, 1) for c in range(43,45)]
    sediment = [ID_sheet.cell_value(c, 2) for c in range(43,45)]
    
    
    diameter = geometory_sheet.cell_value(4, 2)*12
    geometry_len = column_len(geometory_sheet,15)
    geometry_depth = [geometory_sheet.cell_value(c, 2) for c in range(15,geometry_len)]
    # v_N = [geometory_sheet.cell_value(c, 13) for c in range(15,geometry_len)]
    # v_O = [geometory_sheet.cell_value(c, 14) for c in range(15,geometry_len)]
    # v_P = [geometory_sheet.cell_value(c, 15) for c in range(15,geometry_len)]
    # v_Q = [geometory_sheet.cell_value(c, 16) for c in range(15,geometry_len)]
    # v_R = [geometory_sheet.cell_value(c, 17) for c in range(15,geometry_len)]
    # v_S = [geometory_sheet.cell_value(c, 18) for c in range(15,geometry_len)]
    # v_U = [geometory_sheet.cell_value(c, 20) for c in range(15,geometry_len)]
    # v_V = [geometory_sheet.cell_value(c, 21) for c in range(15,geometry_len)]
    # v_W = [geometory_sheet.cell_value(c, 22) for c in range(15,geometry_len)]
    # v_X = [geometory_sheet.cell_value(c, 23) for c in range(15,geometry_len)]

    
    manning_len = column_len(mannings_sheet,1)
    slope = mannings_sheet.cell_value(1, 6)
    slp = 'Slope = ' + str(slope) + ' (ft/ft)'
    manning_d = [mannings_sheet.cell_value(c, 1) for c in range(15,manning_len)]
    manning_18 = [mannings_sheet.cell_value(c, 5) for c in range(15,manning_len)]
    manning_14 = [mannings_sheet.cell_value(c, 6) for c in range(15,manning_len)]
    manning_1 = [mannings_sheet.cell_value(c, 7) for c in range(15,manning_len)]


### Plot QVD and save:
    
    fig, (p1,p2,p3) = plt.subplots(3,1,constrained_layout=True, figsize=(7, 10))
    fig.suptitle('Meter ID: ' + meter_name, y= .92)
    
    p1.plot(meter_x, meter_q, 'black', linewidth=0.5, label='Flow')
    p1.set_xlim(datetime.datetime(2019, 11, 1, 0, 0),datetime.datetime(2020, 4, 15, 0, 0))
    p1.set_ylabel("Flow (mgd)")
    p1.grid(True)
    
    p2.plot(meter_x, meter_d, 'black', linewidth=0.5, label = 'Depth')
    p2.set_xlim(datetime.datetime(2019, 11, 1, 0, 0),datetime.datetime(2020, 4, 15, 0, 0))
    p2.set_ylabel("Depth (in)")
    p2.grid(True)
    
    p3.plot(meter_x, meter_v, 'black', linewidth=0.5, label = 'Depth')
    p3.set_xlim(datetime.datetime(2019, 11, 1, 0, 0),datetime.datetime(2020, 4, 15, 0, 0))
    p3.set_ylabel("Velocity (fps)")
    p3.grid(True)
    
    # Date axis formating:
    plt.gcf().autofmt_xdate()
    myFmt = mdates.DateFormatter('%m-%d-%Y')
    plt.gca().xaxis.set_major_formatter(myFmt)
    
    # plt.show()
   
    png_file = excel_files[a]
    plt.savefig(r"C:\Users\ssin\OneDrive - DHI\Desktop\EBMUD\9_AppendixA_QDV_Scattergraph\Plots\\" +  png_file[:-5] + '_QVD', bbox_inches='tight')
    plt.close(fig)
       
### Plot Scatter Graph

    fig1, ax = plt.subplots(figsize=(7, 9))
    
    ax.scatter(meter_v, meter_d, s=1, c='b', marker ='o', label='Measured Data')
    ax.scatter(meter_v_novdec, meter_d_novdec, s=1, c='r', marker ='o', 
               label='Measured Data (Nov-Dec)')
    ax.plot(manning_18, manning_d, label='Mannings 0.018', linewidth=0.5)
    ax.plot(manning_14, manning_d, label='Mannings 0.014', linewidth=1)
    ax.plot(manning_1, manning_d, label='Mannings 0.01', linewidth=0.5)
    ax.plot(crown_v, crown_d, label='Crown of Pipe', inewidth=2)
    ax.plot(crown_v, sediment, label='Sediment', linewidth=2)
    
    ax.legend(bbox_to_anchor=(-.05, -.18, .7, .08), ncol = 3, loc ='upper left',
              borderaxespad=0)
    ax.set_title('Meter ID: ' + meter_name, y=1.03, pad= 10)
    ax.set_xlabel("Velocity (ft/s)")
    ax.set_ylabel("Depth of Flow (in)")
    ax.grid(True)
     
    # plt.show()
    png_file = excel_files[a]
    plt.savefig(r"C:\Users\ssin\OneDrive - DHI\Desktop\EBMUD\9_AppendixA_QDV_Scattergraph\Plots\\" +  png_file[:-5] + '_scatter', bbox_inches='tight')
    plt.close(fig)                         

### Report - Word document:
    
    Appendix_A = r"C:\Users\ssin\OneDrive - DHI\Desktop\EBMUD\AppendixA_DHI.docx"
    doc = Document(Appendix_A)
        
    # Add plots
    doc.add_picture(r'C:\Users\ssin\OneDrive - DHI\Desktop\EBMUD\9_AppendixA_QDV_Scattergraph\Plots' + '\\' +png_file[:-5] + '_QVD' '.png')
    doc.add_picture(r'C:\Users\ssin\OneDrive - DHI\Desktop\EBMUD\9_AppendixA_QDV_Scattergraph\Plots' + '\\' +png_file[:-5] + '_scatter' +'.png')
               
    doc.save(Appendix_A)
    